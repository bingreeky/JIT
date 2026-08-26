"""OfficeBench tools adapted for local workspace mode."""

from __future__ import annotations

import os
import subprocess
import time
from datetime import datetime
from email import policy
from email.message import EmailMessage
from email.parser import BytesParser
from glob import glob as py_glob
from typing import Any, Callable, Dict, Tuple

from .base import Tool


class _OfficeBenchExecutor:
    def __init__(self):
        self._workspace = ""
        self._testbed_root = "/testbed"
        self._shell_cwd = "/"
        self.current_app = ""

        self._action_handlers: Dict[Tuple[str, str], Callable[[Dict[str, Any]], str]] = {
            ("calendar", "create_event"): self._calendar_create_event,
            ("calendar", "delete_event"): self._calendar_delete_event,
            ("calendar", "list_events"): self._calendar_list_events,
            ("email", "send_email"): self._email_send_email,
            ("email", "list_emails"): self._email_list_emails,
            ("email", "read_email"): self._email_read_email,
            ("excel", "read_file"): self._excel_read_file,
            ("excel", "set_cell"): self._excel_set_cell,
            ("excel", "delete_cell"): self._excel_delete_cell,
            ("excel", "create_new_file"): self._excel_create_new_file,
            ("excel", "convert_to_pdf"): self._excel_convert_to_pdf,
            ("ocr", "recognize_file"): self._ocr_recognize_file,
            ("pdf", "convert_to_image"): self._pdf_convert_to_image,
            ("pdf", "convert_to_word"): self._pdf_convert_to_word,
            ("pdf", "read_file"): self._pdf_read_file,
            ("pdf", "image_convert_to_pdf"): self._pdf_image_convert_to_pdf,
            ("word", "convert_to_pdf"): self._word_convert_to_pdf,
            ("word", "create_new_file"): self._word_create_new_file,
            ("word", "read_file"): self._word_read_file,
            ("word", "write_to_file"): self._word_write_to_file,
            ("llm", "complete_text"): self._llm_complete_text,
            ("system", "copy"): self._system_copy,
            ("system", "paste"): self._system_paste,
        }

    @property
    def workspace(self) -> str:
        return self._workspace

    def set_workspace(self, path: str) -> None:
        self._workspace = os.path.abspath(path)
        self._testbed_root = os.path.join(self._workspace, "testbed")
        os.makedirs(os.path.join(self._testbed_root, "data"), exist_ok=True)
        os.makedirs(os.path.join(self._testbed_root, "emails"), exist_ok=True)
        os.makedirs(os.path.join(self._testbed_root, "calendar"), exist_ok=True)
        self._shell_cwd = self._testbed_root

    def _minor_action_fix(self, args: Dict[str, Any]) -> Dict[str, Any]:
        fixed: Dict[str, Any] = {}
        for k, v in args.items():
            if isinstance(v, list) and len(v) == 1:
                fixed[k] = v[0]
            else:
                fixed[k] = v
        return fixed

    def _map_testbed_path(self, value: Any) -> Any:
        if not isinstance(value, str):
            return value

        aliases = (
            "/testbed",
            "testbed",
            "workspace/testbed",
            "/workspace/testbed",
            "./workspace/testbed",
            "./testbed",
        )
        for alias in aliases:
            if value == alias:
                return self._testbed_root
            if value.startswith(alias + "/"):
                suffix = value[len(alias) + 1 :]
                return os.path.join(self._testbed_root, suffix)

        return value

    def _resolve_user_file_path(self, value: str) -> str:
        value = self._map_testbed_path(value)
        if os.path.isabs(value):
            return value
        return os.path.join(self._testbed_root, value)

    @staticmethod
    def _normalize_path(path: str) -> str:
        return os.path.normpath(path)

    def _run_shell_command(self, command: str) -> str:
        cmd = str(command)

        if cmd.strip().startswith("cd "):
            cd_arg = cmd[cmd.index("cd ") + 3 :].strip()
            cd_arg = str(self._map_testbed_path(cd_arg))
            if os.path.isabs(cd_arg):
                new_path = self._normalize_path(cd_arg)
            else:
                new_path = self._normalize_path(os.path.join(self._shell_cwd, cd_arg))

            if os.path.isdir(new_path):
                self._shell_cwd = new_path
                return "Successfully executed command: cd."
            return f"Failed to execute command: directory not found: {cd_arg}"

        proc = subprocess.run(
            cmd,
            cwd=self._shell_cwd,
            shell=True,
            text=True,
            capture_output=True,
        )

        combined = (proc.stdout or "") + (proc.stderr or "")
        combined = combined.strip()
        if combined:
            return combined
        if proc.returncode == 0:
            return f"Successfully executed command: {cmd}."
        return f"Failed to execute command: {cmd}."

    def _normalize_pathlike_args(self, args: Dict[str, Any]) -> Dict[str, Any]:
        norm_args = dict(args)
        for key, value in list(norm_args.items()):
            if not isinstance(value, str):
                continue
            if key.endswith("_path") or key in {"file_path", "file", "input_file", "output_file"}:
                norm_args[key] = self._resolve_user_file_path(value)
            else:
                norm_args[key] = self._map_testbed_path(value)
        return norm_args

    def _call_action_handler(self, app: str, action: str, args: Dict[str, Any]) -> str:
        handler = self._action_handlers.get((app, action))
        if handler is None:
            return f"Error: unknown OfficeBench action '{app}.{action}'."

        norm_args = self._normalize_pathlike_args(args)
        try:
            return str(handler(norm_args))
        except TypeError as e:
            return f"Error: invalid arguments for '{app}.{action}': {e}"
        except Exception as e:
            return f"Error executing '{app}.{action}': {e}"

    @staticmethod
    def _format_calendar_time(obj):
        if obj:
            return obj.dt.strftime("%Y-%m-%d %H:%M:%S")
        return obj

    def _calendar_create_event(self, args: Dict[str, Any]) -> str:
        user = args["user"]
        summary = args["summary"]
        time_start = args["time_start"]
        time_end = args["time_end"]
        if user == "Multiple users":
            return f"OBSERVATION: Failed to create a new event to {user}. Only support one user."

        from icalendar import Calendar, Event

        os.makedirs(os.path.join(self._testbed_root, "calendar"), exist_ok=True)
        try:
            calendar_file = os.path.join(self._testbed_root, "calendar", f"{user}.ics")
            if not os.path.exists(calendar_file):
                calendar = Calendar()
                calendar.add("prodid", "-//My Calendar Product//mxm.dk//")
                calendar.add("version", "2.0")
            else:
                with open(calendar_file, "rb") as f:
                    calendar = Calendar.from_ical(f.read())

            event = Event()
            event.add("summary", summary)
            event.add("dtstart", datetime.strptime(time_start, "%Y-%m-%d %H:%M:%S"))
            event.add("dtend", datetime.strptime(time_end, "%Y-%m-%d %H:%M:%S"))
            event.add("dtstamp", datetime.now())
            event.add("description", "This is a test event")
            event.add("location", "Online")
            calendar.add_component(event)

            with open(calendar_file, "wb") as f:
                f.write(calendar.to_ical())
            return f"OBSERVATION: Successfully create a new event to {user}'s calendar."
        except Exception:
            return f"OBSERVATION: Failed to create a new event to {user}'s calendar."

    def _calendar_delete_event(self, args: Dict[str, Any]) -> str:
        user = args["user"]
        summary = args["summary"]

        from icalendar import Calendar

        try:
            calendar_file = os.path.join(self._testbed_root, "calendar", f"{user}.ics")
            if not os.path.exists(calendar_file):
                return f"OBSERVATION: Failed to delete an event named {summary} from {user}'s calendar."
            with open(calendar_file, "rb") as f:
                calendar = Calendar.from_ical(f.read())

            to_remove = None
            for component in calendar.walk():
                if component.name == "VEVENT" and component.get("summary") == summary:
                    to_remove = component
                    break
            if to_remove is not None:
                calendar.subcomponents.remove(to_remove)

            with open(calendar_file, "wb") as f:
                f.write(calendar.to_ical())
            return f"OBSERVATION: Successfully delete an event named {summary} from {user}'s calendar."
        except Exception:
            return f"OBSERVATION: Failed to delete an event named {summary} from {user}'s calendar."

    def _calendar_list_events(self, args: Dict[str, Any]) -> str:
        username = args["username"]
        if username == "Multiple users":
            return f"OBSERVATION: Failed to list events for {username}. Only support one user."

        import icalendar

        calendar_file = os.path.join(self._testbed_root, "calendar", f"{username}.ics")
        if not os.path.exists(calendar_file):
            os.makedirs(os.path.dirname(calendar_file), exist_ok=True)
            calendar = icalendar.Calendar()
            calendar.add("prodid", "-//My Calendar Product//mxm.dk//")
            calendar.add("version", "2.0")
            with open(calendar_file, "wb") as f:
                f.write(calendar.to_ical())

        try:
            with open(calendar_file, "rb") as f:
                calendar = icalendar.Calendar.from_ical(f.read())
            message = ""
            for component in calendar.walk():
                if component.name == "VEVENT":
                    message += f'Summary: {component.get("summary")}\n'
                    message += f"Start Time: {self._format_calendar_time(component.get('dtstart'))}\n"
                    message += f"End Time: {self._format_calendar_time(component.get('dtend'))}\n"
                    message += f"Description: {component.get('description')}\n"
                    message += f"Location: {component.get('location')}\n"
                    message += "-" * 50 + "\n"
            message = message.strip()
        except Exception:
            message = "Error: Failed to list events."

        if message == "Error: Failed to list events.":
            return f"OBSERVATION: Failed to list events for {username}."
        if message == "":
            return f"OBSERVATION: No events found for {username}."
        return f"OBSERVATION: Successfully list events for {username}:\n{message}"

    @staticmethod
    def _get_email_content(msg):
        if msg.is_multipart():
            parts = []
            for part in msg.iter_parts():
                if part.get_content_type() in {"text/plain", "text/html"}:
                    charset = part.get_content_charset() or "utf-8"
                    parts.append(part.get_payload(decode=True).decode(charset, errors="replace"))
            return "\n".join(parts)
        charset = msg.get_content_charset() or "utf-8"
        payload = msg.get_payload(decode=True)
        if payload is None:
            return ""
        return payload.decode(charset, errors="replace")

    def _email_send_email(self, args: Dict[str, Any]) -> str:
        sender = str(args["sender"])
        recipient = str(args["recipient"])
        subject = str(args["subject"])
        content = str(args["content"])

        if recipient == "Multiple recipients":
            return f"OBSERVATION: Failed to send email to {recipient}. Only support one recipient."

        if "@" in sender:
            sender = sender.split("@")[0]
        if "@" in recipient:
            recipient = recipient.split("@")[0]

        os.makedirs(os.path.join(self._testbed_root, "emails", sender), exist_ok=True)
        os.makedirs(os.path.join(self._testbed_root, "emails", recipient), exist_ok=True)
        try:
            email = EmailMessage()
            email["From"] = sender + "@example.com"
            email["To"] = recipient + "@example.com"
            email["Subject"] = subject
            email.set_content(content)

            recipient_file = os.path.join(self._testbed_root, "emails", recipient, f"{subject}.eml")
            with open(recipient_file, "w", encoding="utf-8") as f:
                f.write(email.as_string())

            sender_file = os.path.join(self._testbed_root, "emails", sender, f"{subject}.eml")
            with open(sender_file, "w", encoding="utf-8") as f:
                f.write(email.as_string())

            return f"OBSERVATION: Successfully sent email to {recipient}."
        except Exception:
            return f"OBSERVATION: Failed to send email to {recipient}."

    def _email_list_emails(self, args: Dict[str, Any]) -> str:
        username = str(args["username"])
        os.makedirs(os.path.join(self._testbed_root, "emails", username), exist_ok=True)
        try:
            email_folder = os.path.join(self._testbed_root, "emails", username)
            email_files = py_glob(os.path.join(email_folder, "*.eml"))
            message = ""
            for email_file in email_files:
                with open(email_file, "rb") as f:
                    email_content = f.read()
                    email_msg = BytesParser(policy=policy.default).parsebytes(email_content)
                email_name = os.path.basename(email_file)
                message += f"Email ID: {email_name}\n"
                message += f"From: {email_msg['From']}\n"
                message += f"To: {email_msg['To']}\n"
                message += f"Subject: {email_msg['Subject']}\n"
                message += f"Content: {self._get_email_content(email_msg)[:20] + '...'}\n"
                message += "-" * 50 + "\n"

            if not message:
                return f"OBSERVATION: No emails found for {username}."
            return f"OBSERVATION: Successfully list emails for {username}:\n{message}"
        except Exception:
            return f"OBSERVATION: Failed to list emails for {username}."

    def _email_read_email(self, args: Dict[str, Any]) -> str:
        username = str(args["username"])
        email_id = str(args["email_id"])

        os.makedirs(os.path.join(self._testbed_root, "emails", username), exist_ok=True)
        try:
            if not email_id.endswith(".eml"):
                email_id += ".eml"
            email_file = os.path.join(self._testbed_root, "emails", username, email_id)
            with open(email_file, "rb") as f:
                email_content = f.read()
                email_msg = BytesParser(policy=policy.default).parsebytes(email_content)

            message = ""
            message += f"From: {email_msg['From']}\n"
            message += f"To: {email_msg['To']}\n"
            message += f"Subject: {email_msg['Subject']}\n"
            message += f"Content: {self._get_email_content(email_msg) + '...'}\n"
            return f"OBSERVATION: Successfully read email {email_id} for {username}:\n{message}"
        except Exception:
            return f"OBSERVATION: Failed to read email {email_id} for {username}."

    def _excel_read_file(self, args: Dict[str, Any]) -> str:
        file_path = args["file_path"]
        sheet = args.get("sheet")

        import openpyxl

        if not os.path.exists(file_path):
            return f"OBSERVATION: The file {file_path} does not exist. Failed to read the file."

        if sheet is None:
            ws = openpyxl.load_workbook(file_path).active
        else:
            ws = openpyxl.load_workbook(file_path)[sheet]

        content_string = ""
        for row in ws.iter_rows():
            for cell in row:
                row_idx = cell.row
                col_idx = cell.column
                value = cell.value if cell.value is not None else "[Empty Cell]"
                content_string += f"({row_idx}, {col_idx}): {value}\t"
            content_string += "\n"

        return f"OBSERVATION: The following is the table from the excel file:\n{content_string}"

    def _excel_set_cell(self, args: Dict[str, Any]) -> str:
        file_path = args["file_path"]
        text = args["text"]
        row_idx = args["row_idx"]
        column_idx = args["column_idx"]
        sheet_name = args.get("sheet_name")

        import openpyxl

        if not os.path.exists(file_path):
            return f"OBSERVATION: The file {file_path} does not exist. Failed to write to the file."

        if text is True:
            text = ""
        try:
            workbook = openpyxl.load_workbook(file_path)
            if sheet_name is None:
                sheet = workbook.active
            else:
                try:
                    sheet = workbook[sheet_name]
                except KeyError:
                    sheet = workbook.create_sheet(title=sheet_name)

            sheet.cell(row=int(row_idx), column=int(column_idx), value=text)
            workbook.save(file_path)
            return f"OBSERVATION: Successfully write text to {file_path}"
        except Exception:
            return f"OBSERVATION: Failed to write text to {file_path}"

    def _excel_delete_cell(self, args: Dict[str, Any]) -> str:
        file_path = args["file_path"]
        row_idx = args["row_idx"]
        column_idx = args["column_idx"]
        sheet_name = args.get("sheet_name")

        import openpyxl

        if not os.path.exists(file_path):
            return f"OBSERVATION: The file {file_path} does not exist. Failed to delete the cell in the file."

        try:
            workbook = openpyxl.load_workbook(file_path)
            if sheet_name is None:
                sheet = workbook.active
            else:
                try:
                    sheet = workbook[sheet_name]
                except KeyError:
                    sheet = workbook.create_sheet(title=sheet_name)

            sheet.cell(row=int(row_idx), column=int(column_idx), value="")
            workbook.save(file_path)
            return f"OBSERVATION: Successfully delete a cell in {file_path}"
        except Exception:
            return f"OBSERVATION: Failed to delete a cell in {file_path}"

    def _excel_create_new_file(self, args: Dict[str, Any]) -> str:
        file_path = args["file_path"]

        import openpyxl

        if os.path.exists(file_path):
            return f"OBSERVATION: File {file_path} already exists"

        try:
            wb = openpyxl.Workbook()
            wb.save(file_path)
            return f"OBSERVATION: Successfully create new file {file_path}"
        except Exception:
            return f"OBSERVATION: Failed to create new file {file_path}"

    def _excel_convert_to_pdf(self, args: Dict[str, Any]) -> str:
        excel_file_path = args["excel_file_path"]
        pdf_file_path = args["pdf_file_path"]

        if not os.path.exists(excel_file_path):
            return f"OBSERVATION: {excel_file_path} does not exist. Failed to convert {excel_file_path} to {pdf_file_path}"

        try:
            output_dir = os.path.dirname(pdf_file_path)
            subprocess.call(["libreoffice", "--headless", "--convert-to", "pdf", excel_file_path, "--outdir", output_dir])
            return f"OBSERVATION: Successfully convert {excel_file_path} to {pdf_file_path}"
        except Exception:
            return f"OBSERVATION: Failed to convert {excel_file_path} to {pdf_file_path}"

    def _ocr_recognize_file(self, args: Dict[str, Any]) -> str:
        file_path = args["file_path"]

        if not os.path.exists(file_path):
            return f"OBSERVATION: The file {file_path} does not exist. Failed to recognize text."

        try:
            import pytesseract
            from PIL import Image

            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)
        except Exception:
            text = None

        if text:
            return f"OBSERVATION: The text from {file_path} is:\n{text}"
        return f"OBSERVATION: Failed to recognize text from {file_path}"

    def _pdf_convert_to_image(self, args: Dict[str, Any]) -> str:
        pdf_file_path = args["pdf_file_path"]
        image_file_path = args["image_file_path"]

        if not os.path.exists(pdf_file_path):
            return f"OBSERVATION: The file {pdf_file_path} does not exist. Failed to convert pdf to image."

        try:
            import fitz

            with fitz.open(pdf_file_path) as doc:
                page = doc.load_page(0)
                pix = page.get_pixmap()
                pix.save(image_file_path)
            status = "Success"
        except Exception as e:
            status = f"Failed: {str(e)}"

        return "OBSERVATION: " + status

    def _pdf_convert_to_word(self, args: Dict[str, Any]) -> str:
        pdf_file_path = args["pdf_file_path"]
        word_file_path = args["word_file_path"]

        if not os.path.exists(pdf_file_path):
            return f"OBSERVATION: The pdf file {pdf_file_path} does not exist. Failed to convert the file to word."

        try:
            from pdf2docx import Converter

            cv = Converter(pdf_file_path)
            cv.convert(word_file_path, start=0, end=None)
            cv.close()
            status = "Success"
        except Exception as e:
            status = f"Failed: {str(e)}"

        return "OBSERVATION: " + status

    def _pdf_read_file(self, args: Dict[str, Any]) -> str:
        pdf_file_path = args["pdf_file_path"]

        if not os.path.exists(pdf_file_path):
            return f"OBSERVATION: The pdf file {pdf_file_path} does not exist. Failed to read the file."

        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(pdf_file_path)
            pages = reader.pages
            text = ""
            for page in pages:
                text += page.extract_text()
            return "OBSERVATION: " + text
        except Exception as e:
            return f"OBSERVATION: Failed: {e}"

    def _pdf_image_convert_to_pdf(self, args: Dict[str, Any]) -> str:
        image_file_path = args["image_file_path"]
        pdf_file_path = args["pdf_file_path"]

        if not os.path.exists(image_file_path):
            return f"OBSERVATION: The file {image_file_path} does not exist. Failed to convert image to pdf."

        try:
            from PIL import Image

            with Image.open(image_file_path) as img:
                if img.mode != "RGB":
                    img = img.convert("RGB")
                img.save(pdf_file_path, "PDF", resolution=100.0)
            status = "Success"
        except Exception as e:
            status = f"Failed: {str(e)}"

        return "OBSERVATION: " + status

    def _word_convert_to_pdf(self, args: Dict[str, Any]) -> str:
        word_file_path = args["word_file_path"]
        pdf_file_path = args["pdf_file_path"]

        if not os.path.exists(word_file_path):
            return f"OBSERVATION: The word file {word_file_path} does not exist. Failed to convert the file to pdf."

        try:
            output_dir = os.path.dirname(pdf_file_path)
            subprocess.call(["libreoffice", "--headless", "--convert-to", "pdf", word_file_path, "--outdir", output_dir])
            success = True
        except Exception:
            success = False

        if success:
            return f"OBSERVATION: Successfully convert {word_file_path} to {pdf_file_path}"
        return f"OBSERVATION: Failed to convert {word_file_path} to {pdf_file_path}"

    def _word_create_new_file(self, args: Dict[str, Any]) -> str:
        file_path = args["file_path"]

        try:
            from docx import Document

            document = Document()
            document.save(file_path)
            return f"OBSERVATION: Successfully create new file {file_path}"
        except Exception:
            return f"OBSERVATION: Failed to create new file {file_path}"

    def _word_read_file(self, args: Dict[str, Any]) -> str:
        file_path = args["file_path"]

        if not os.path.exists(file_path):
            return f"OBSERVATION: The file {file_path} does not exist. Failed to read the file."

        try:
            from docx import Document

            doc = Document(file_path)
            observation = "The following is the content from the word file:"
            for paragraph in doc.paragraphs:
                observation += f"\n{paragraph.text}"
            return "OBSERVATION: " + observation
        except Exception as e:
            return f"OBSERVATION: Failed to read the file. {e}"

    def _word_write_to_file(self, args: Dict[str, Any]) -> str:
        file_path = args["file_path"]
        contents = args["contents"]
        style = args.get("style", "pure-text")

        if not os.path.exists(file_path):
            return f"OBSERVATION: The file {file_path} does not exist. Failed to write to the file."

        try:
            from docx import Document

            document = Document(file_path)
            if style == "pure-text":
                document.add_paragraph(contents)
            elif style == "title":
                document.add_heading(contents, 0)
            elif style == "subtitle":
                document.add_heading(contents, 1)
            else:
                raise NotImplementedError
            document.save(file_path)
            return f"OBSERVATION: Successfully write contents to {file_path}"
        except Exception:
            return f"OBSERVATION: Failed to write contents to {file_path}"

    def _llm_complete_text(self, args: Dict[str, Any]) -> str:
        prompt = str(args["prompt"])
        try:
            import openai

            model_name = "gpt-5.3-codex"
            with open("/openai_key.txt", "r", encoding="utf-8") as f:
                key = f.read().strip()

            messages = [
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt},
            ]

            gpt_responses = None
            retry_num = 0
            retry_limit = 2
            error = None
            while gpt_responses is None and retry_num < retry_limit:
                try:
                    gpt_responses = openai.ChatCompletion.create(
                        model=model_name,
                        messages=messages,
                        stop=None,
                        api_key=key,
                        temperature=0,
                        n=1,
                        top_p=1,
                        max_tokens=50,
                    )
                    error = None
                except Exception as e:
                    error = e
                    time.sleep(5)
                    retry_num += 1
            if error:
                raise Exception(error)
            return gpt_responses["choices"][0]["message"]["content"]
        except Exception as e:
            return f"Error: {e}"

    def _clipboard_path(self) -> str:
        # Per-testbed clipboard: /tmp/.clipboard is a GLOBAL file that leaks
        # text between concurrent (case, rollout) units and between parallel
        # eval processes; the testbed root is unique per unit.
        return os.path.join(self._testbed_root, ".clipboard")

    def _system_copy(self, args: Dict[str, Any]) -> str:
        text = str(args.get("text", ""))
        try:
            with open(self._clipboard_path(), "w", encoding="utf-8") as f:
                f.write(text)
            return "Text copied successfully."
        except Exception:
            return "Failed to copy the text."

    def _system_paste(self, args: Dict[str, Any]) -> str:
        _ = args
        try:
            with open(self._clipboard_path(), "r", encoding="utf-8") as f:
                return f.read()
        except Exception:
            return "Failed to paste."

    def execute(self, app: str, action: str, args: Dict[str, Any]) -> str:
        app = str(app)
        action = str(action)
        args = self._minor_action_fix(args or {})

        if action == "switch_app":
            app = "system"

        if app == "system":
            if action == "switch_app":
                target = str(args.get("target_app", "")).strip()
                if not target:
                    return "Error: 'target_app' is required for switch_app."
                self.current_app = target
                return f"Successfully switched to app: {target}"
            if action == "finish_task":
                answer = str(args.get("answer", "None"))
                answer_path = os.path.join(self._testbed_root, "data", "answer.txt")
                os.makedirs(os.path.dirname(answer_path), exist_ok=True)
                with open(answer_path, "w", encoding="utf-8") as f:
                    f.write(answer)
                return "Task finished"
            if action == "got_stuck":
                return "Task failed"
            return self._call_action_handler(app, action, args)

        if app == "shell" and action == "command":
            command = args.get("command", "")
            if isinstance(command, list):
                command = " ".join(str(x) for x in command)
            return self._run_shell_command(str(command))

        return self._call_action_handler(app, action, args)


class OfficeBenchActionTool(Tool):
    name = "officebench_action"
    description = (
        "Execute a single OfficeBench action in local workspace mode. "
        "Call format: {app: string, action: string, args: dict}. "
        "Relative file paths inside args are resolved under /testbed; absolute /testbed/... paths are remapped to the current task workspace. "
        "Do not call imaginary standalone tools such as `shell.command`; use this tool with app=`shell`, action=`command`, and args={command: ...}. "
        "The tool always returns a single string. Most actions return an 'OBSERVATION: ...' message; "
        "system.copy returns a clipboard status string, system.paste returns pasted text or an error string, "
        "shell.command returns stdout/stderr or a success/failure message, and llm.complete_text returns raw model text or an error string.\n\n"
        "Supported actions and input/output formats:\n"
        "- calendar.create_event: args={user: str, summary: str, time_start: str('%Y-%m-%d %H:%M:%S'), time_end: str('%Y-%m-%d %H:%M:%S')}. Output: string observation indicating whether the event was created.\n"
        "- calendar.delete_event: args={user: str, summary: str}. Output: string observation indicating whether the matching event was deleted.\n"
        "- calendar.list_events: args={username: str}. Output: string observation containing all events for that user, or a no-events / failure message.\n"
        "- email.send_email: args={sender: str, recipient: str, subject: str, content: str}. Output: string observation indicating whether the email was sent. Only one recipient is supported.\n"
        "- email.list_emails: args={username: str}. Output: string observation listing email id/from/to/subject/content preview for that user, or a no-emails / failure message.\n"
        "- email.read_email: args={username: str, email_id: str}. Output: string observation containing full email fields and content for that email id.\n"
        "- excel.read_file: args={file_path: str, sheet: str | optional}. Output: string observation containing sheet contents as '(row, col): value' cells.\n"
        "- excel.set_cell: args={file_path: str, row_idx: int|str, column_idx: int|str, text: Any, sheet_name: str | optional}. Output: string observation indicating whether the target cell was written.\n"
        "- excel.delete_cell: args={file_path: str, row_idx: int|str, column_idx: int|str, sheet_name: str | optional}. Output: string observation indicating whether the target cell was cleared.\n"
        "- excel.create_new_file: args={file_path: str}. Output: string observation indicating whether a new workbook was created.\n"
        "- excel.convert_to_pdf: args={excel_file_path: str, pdf_file_path: str}. Output: string observation indicating whether conversion succeeded.\n"
        "- ocr.recognize_file: args={file_path: str}. Output: string observation containing recognized text, or a failure message.\n"
        "- pdf.convert_to_image: args={pdf_file_path: str, image_file_path: str}. Output: string observation with conversion status.\n"
        "- pdf.convert_to_word: args={pdf_file_path: str, word_file_path: str}. Output: string observation with conversion status.\n"
        "- pdf.read_file: args={pdf_file_path: str}. Output: string observation containing extracted text from all pages.\n"
        "- pdf.image_convert_to_pdf: args={image_file_path: str, pdf_file_path: str}. Output: string observation with conversion status.\n"
        "- shell.command: args={command: str | list[str]}. Output: command stdout/stderr if any; otherwise a success/failure string. A leading 'cd ...' updates the shell working directory for later shell actions. Prefer relative paths like `data`, `emails`, and `calendar` in shell commands.\n"
        "- word.convert_to_pdf: args={word_file_path: str, pdf_file_path: str}. Output: string observation indicating whether conversion succeeded.\n"
        "- word.create_new_file: args={file_path: str}. Output: string observation indicating whether a new document was created.\n"
        "- word.read_file: args={file_path: str}. Output: string observation containing the document paragraphs.\n"
        "- word.write_to_file: args={file_path: str, contents: str, style: str | optional('pure-text'|'title'|'subtitle')}. Output: string observation indicating whether content was appended.\n"
        "- llm.complete_text: args={prompt: str}. Output: raw completion text, or an 'Error: ...' string.\n"
        "- system.copy: args={text: str}. Output: clipboard status string such as 'Text copied successfully.' or failure.\n"
        "- system.paste: args={}. Output: pasted clipboard text, or 'Failed to paste.'.\n"
        "- system.switch_app: args={target_app: str}. Output: success string indicating the active app name.\n"
        "- system.finish_task: args={answer: str}. Output: 'Task finished' and writes the answer to `data/answer.txt` inside the current task testbed.\n"
        "- system.got_stuck: args={}. Output: 'Task failed'."
    )
    inputs = {
        "app": {
            "type": "string",
            "description": "OfficeBench app name, e.g. system/excel/word/pdf/email/calendar/shell/ocr/llm.",
        },
        "action": {
            "type": "string",
            "description": "Action name under the app.",
        },
        "args": {
            "type": "dict",
            "description": "Action arguments as a dictionary.",
        },
    }
    output_type = "string"

    def __init__(self, model=None):
        super().__init__()
        self._executor = _OfficeBenchExecutor()

    @property
    def workspace(self) -> str:
        return self._executor.workspace

    def set_workspace(self, path: str) -> None:
        self._executor.set_workspace(path)

    def forward(self, app: str, action: str, args: Dict[str, Any]) -> str:
        if not isinstance(args, dict):
            return "Error: 'args' must be a dict."
        return self._executor.execute(app=app, action=action, args=args)


__all__ = [
    "OfficeBenchActionTool",
]
